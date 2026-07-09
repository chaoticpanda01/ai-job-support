"""
Resume text extraction.

Supports:
  application/pdf  — pypdf2 (PdfReader)
  application/vnd.openxmlformats-officedocument.wordprocessingml.document — python-docx

Returns cleaned plain text. All extraction is synchronous (CPU-bound, but fast
for typical resume sizes). Call from a thread pool executor if needed inside
async context (see note in analyze endpoint).
"""

from __future__ import annotations

import io
import logging
import re

logger = logging.getLogger(__name__)

MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_MAX_CHARS = 20_000  # truncate before sending to Claude


class ParseError(Exception):
    """Raised when text extraction fails."""


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    """
    Extract and clean plain text from a PDF or DOCX byte payload.
    Raises ParseError on failure or unsupported MIME type.
    """
    if mime_type == MIME_PDF:
        raw = _extract_pdf(file_bytes)
    elif mime_type == MIME_DOCX:
        raw = _extract_docx(file_bytes)
    else:
        raise ParseError(f"Unsupported MIME type: {mime_type}")

    cleaned = _clean(raw)
    if not cleaned.strip():
        raise ParseError("No readable text found in file")

    if len(cleaned) > _MAX_CHARS:
        logger.warning("Resume text truncated: original=%d chars", len(cleaned))
        cleaned = cleaned[:_MAX_CHARS]

    return cleaned


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import PyPDF2
    except ImportError as exc:
        raise ParseError("PyPDF2 is not installed") from exc

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            parts.append(text)
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        raise ParseError(f"PDF extraction failed: {exc}") from exc


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ParseError("python-docx is not installed") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = [para.text for para in doc.paragraphs]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        raise ParseError(f"DOCX extraction failed: {exc}") from exc


# ---------------------------------------------------------------------------
# Cleaner
# ---------------------------------------------------------------------------


def _clean(text: str) -> str:
    """Normalize whitespace and remove non-printable characters."""
    # Remove null bytes and control characters (keep newlines/tabs)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Collapse runs of blank lines to a single blank line
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse horizontal whitespace (spaces/tabs) within a line
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()
